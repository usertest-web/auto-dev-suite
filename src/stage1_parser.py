# src/stage1_parser.py
from dataclasses import dataclass, field
from typing import Optional
import json
from docx import Document


@dataclass
class PageSettings:
    size: str = "A4"
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5
    margin_left_cm: float = 3.0
    margin_right_cm: float = 2.5

    def to_dict(self) -> dict:
        return {
            "size": self.size,
            "margin_top_cm": self.margin_top_cm,
            "margin_bottom_cm": self.margin_bottom_cm,
            "margin_left_cm": self.margin_left_cm,
            "margin_right_cm": self.margin_right_cm,
        }


@dataclass
class ParsedTemplate:
    page: PageSettings = field(default_factory=PageSettings)
    raw_styles: dict = field(default_factory=dict)
    structure: dict = field(default_factory=dict)
    declaration_text: str = ""

    def to_dict(self) -> dict:
        return {
            "page": self.page.to_dict(),
            "styles": {},
            "raw_styles": self.raw_styles,
            "structure": self.structure,
            "declaration_text": self.declaration_text,
            "citation_format": "GB/T 7714-2015",
            "citation_style": "numbered",
        }


def parse_template(template_path: str) -> ParsedTemplate:
    doc = Document(template_path)
    section = doc.sections[0]
    page = PageSettings(
        margin_top_cm=round(section.top_margin.cm, 1),
        margin_bottom_cm=round(section.bottom_margin.cm, 1),
        margin_left_cm=round(section.left_margin.cm, 1),
        margin_right_cm=round(section.right_margin.cm, 1),
    )

    raw_styles = {}
    for style in doc.styles:
        if style.type is not None:
            style_info = {}
            if style.font.name:
                style_info["font"] = style.font.name
            if style.font.size:
                style_info["size_pt"] = style.font.size.pt
            style_info["bold"] = style.font.bold
            if style.paragraph_format.alignment is not None:
                style_info["alignment"] = style.paragraph_format.alignment
            if style.paragraph_format.space_before:
                style_info["space_before_pt"] = style.paragraph_format.space_before.pt
            if style.paragraph_format.space_after:
                style_info["space_after_pt"] = style.paragraph_format.space_after.pt
            if style.paragraph_format.line_spacing:
                style_info["line_spacing"] = style.paragraph_format.line_spacing
            if style.paragraph_format.first_line_indent:
                style_info["first_indent_cm"] = round(style.paragraph_format.first_line_indent.cm, 2)
            raw_styles[style.name] = style_info

    structure = {
        "cover_fields": ["title", "author", "student_id", "advisor", "date"],
        "chapter_sequence": [
            "cover", "declaration", "abstract_cn", "abstract_en",
            "toc", "chapters", "references", "acknowledgement"
        ],
    }

    return ParsedTemplate(
        page=page,
        raw_styles=raw_styles,
        structure=structure,
    )


def parse_format_spec(spec_path: str, llm_client) -> dict:
    """Use LLM to extract formatting rules from a natural-language spec document."""
    try:
        doc = Document(spec_path)
        spec_text = "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        with open(spec_path, "r", encoding="utf-8") as f:
            spec_text = f.read()

    system = "你是论文排版格式专家。从格式规范文档中提取结构化的排版规则。"
    prompt = f"""从以下格式规范文档中提取所有排版规则，以 JSON 格式返回。

返回格式示例：
{{
  "styles": {{
    "cover_title": {{"font": "黑体", "size_pt": 22, "bold": true, "align": "center"}},
    "heading_1": {{"font": "黑体", "size_pt": 16, "bold": true}},
    "heading_2": {{"font": "黑体", "size_pt": 14, "bold": true}},
    "body": {{"font": "宋体", "size_pt": 12, "line_spacing": 1.5, "first_indent_chars": 2}},
    "caption": {{"font": "宋体", "size_pt": 10, "align": "center"}},
    "reference": {{"font": "宋体", "size_pt": 10.5, "hanging_indent_chars": 2}}
  }},
  "header_footer": {{"header_content": "XX大学本科毕业论文", "page_number_position": "bottom_center"}},
  "special_rules": ["章标题另起一页", "图表编号按章重置"]
}}

规范文档内容：
{spec_text}

请只返回 JSON，不要包含其他解释文字。"""

    response = llm_client.complete(system=system, prompt=prompt, temperature=0.1)
    return json.loads(response.text)


def merge_config(template: ParsedTemplate, spec_rules: dict) -> dict:
    """Merge python-docx extracted data with LLM-parsed spec rules. Spec rules take precedence."""
    result = template.to_dict()
    if "styles" in spec_rules:
        result["styles"] = spec_rules["styles"]
    if "header_footer" in spec_rules:
        result["header_footer"] = spec_rules["header_footer"]
    if "special_rules" in spec_rules:
        result["special_rules"] = spec_rules["special_rules"]
    return result

from dataclasses import dataclass, field


@dataclass
class Issue:
    severity: str  # "error", "warning", "info"
    dimension: str
    detail: str
    location: str = ""
    auto_fixable: bool = False

    def to_dict(self) -> dict:
        return {
            "severity": self.severity,
            "dimension": self.dimension,
            "detail": self.detail,
            "location": self.location,
            "auto_fixable": self.auto_fixable,
        }


@dataclass
class VerificationReport:
    total_checks: int = 0
    passed: int = 0
    failed: int = 0
    issues: list = field(default_factory=list)
    auto_fixed_count: int = 0

    @property
    def pass_(self) -> bool:
        return self.failed == 0

    def to_dict(self) -> dict:
        return {
            "pass": self.pass_,
            "total_checks": self.total_checks,
            "passed": self.passed,
            "failed": self.failed,
            "issues": [i.to_dict() for i in self.issues],
            "auto_fixed_count": self.auto_fixed_count,
        }


def _check_page_settings(doc_path: str, config: dict) -> list:
    issues = []
    try:
        from docx import Document
        doc = Document(doc_path)
        section = doc.sections[0]

        page_cfg = config.get("page", {})
        expected_top = page_cfg.get("margin_top_cm", 2.5)
        expected_bottom = page_cfg.get("margin_bottom_cm", 2.5)
        expected_left = page_cfg.get("margin_left_cm", 3.0)
        expected_right = page_cfg.get("margin_right_cm", 2.5)

        actual_top = round(section.top_margin.cm, 1)
        if abs(actual_top - expected_top) > 0.2:
            issues.append(Issue(
                severity="error", dimension="页面级",
                detail=f"上边距为 {actual_top}cm，要求 {expected_top}cm",
                location="全文", auto_fixable=True,
            ))

        actual_left = round(section.left_margin.cm, 1)
        if abs(actual_left - expected_left) > 0.2:
            issues.append(Issue(
                severity="error", dimension="页面级",
                detail=f"左边距为 {actual_left}cm，要求 {expected_left}cm",
                location="全文", auto_fixable=True,
            ))
    except Exception as e:
        issues.append(Issue(
            severity="error", dimension="页面级",
            detail=f"无法读取页面设置：{e}", auto_fixable=False,
        ))
    return issues


def _check_structure(doc_path: str, config: dict) -> list:
    issues = []
    try:
        from docx import Document
        doc = Document(doc_path)

        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        heading_texts = [h.text.strip() for h in headings]

        required_elements = ["绪论", "参考文献", "致谢"]
        for elem in required_elements:
            found = any(elem in h for h in heading_texts)
            if not found:
                issues.append(Issue(
                    severity="warning", dimension="结构完整性",
                    detail=f"未发现'{elem}'部分",
                    auto_fixable=False,
                ))
    except Exception as e:
        issues.append(Issue(
            severity="error", dimension="结构完整性",
            detail=f"无法读取文档结构：{e}", auto_fixable=False,
        ))
    return issues


def _check_citation_consistency(doc_path: str, config: dict) -> list:
    issues = []
    try:
        from docx import Document
        import re
        doc = Document(doc_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        citation_pattern = re.findall(r'\[([0-9,\-\s]+)\]', full_text)
        cited_numbers = set()
        for match in citation_pattern:
            for part in re.findall(r'\d+', match):
                cited_numbers.add(int(part))

        if cited_numbers:
            max_cited = max(cited_numbers)
            expected = set(range(1, max_cited + 1))
            missing = sorted(expected - cited_numbers)
            if missing:
                issues.append(Issue(
                    severity="warning", dimension="引用一致性",
                    detail=f"引用编号不连续，缺少：{missing}",
                    auto_fixable=False,
                ))

            jumps = []
            prev = 0
            for n in sorted(cited_numbers):
                if n - prev > 5:
                    jumps.append(f"{prev}->{n}")
                prev = n
            if jumps:
                issues.append(Issue(
                    severity="info", dimension="引用一致性",
                    detail=f"引用编号跳号：{', '.join(jumps)}",
                    auto_fixable=False,
                ))
    except Exception as e:
        issues.append(Issue(
            severity="error", dimension="引用一致性",
            detail=f"无法检查引用：{e}", auto_fixable=False,
        ))
    return issues


def verify_format(doc_path: str, template_config: dict) -> VerificationReport:
    issues = []
    issues.extend(_check_page_settings(doc_path, template_config))
    issues.extend(_check_structure(doc_path, template_config))
    issues.extend(_check_citation_consistency(doc_path, template_config))

    total = max(len(issues), 1)
    errors_and_warnings = [i for i in issues if i.severity in ("error", "warning")]
    passed = max(0, total - len(errors_and_warnings))
    failed = len(errors_and_warnings)

    return VerificationReport(
        total_checks=total,
        passed=passed,
        failed=failed,
        issues=issues,
    )
